from langchain_community.document_loaders import YoutubeLoader,WebBaseLoader
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
load_dotenv()
from IPython.display import Markdown, display
from langchain.text_splitter import RecursiveCharacterTextSplitter
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup
from nltk.tokenize import word_tokenize
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.summarize import load_summarize_chain
from langchain.chains.question_answering import load_qa_chain
models=ChatGoogleGenerativeAI(model='gemini-1.5-pro')
languages = {
    "ab": "Abkhazian",
    "aa": "Afar",
    "af": "Afrikaans",
    "ak": "Akan",
    "sq": "Albanian",
    "am": "Amharic",
    "ar": "Arabic",
    "hy": "Armenian",
    "as": "Assamese",
    "ay": "Aymara",
    "az": "Azerbaijani",
    "bn": "Bangla",
    "ba": "Bashkir",
    "eu": "Basque",
    "be": "Belarusian",
    "bho": "Bhojpuri",
    "bs": "Bosnian",
    "br": "Breton",
    "bg": "Bulgarian",
    "my": "Burmese",
    "ca": "Catalan",
    "ceb": "Cebuano",
    "zh-Hans": "Chinese (Simplified)",
    "zh-Hant": "Chinese (Traditional)",
    "co": "Corsican",
    "hr": "Croatian",
    "cs": "Czech",
    "da": "Danish",
    "dv": "Divehi",
    "nl": "Dutch",
    "dz": "Dzongkha",
    "en": "English",
     'en-US' : 'English (United States)',
     'en-IN':'English(Indian)',
    "eo": "Esperanto",
    "et": "Estonian",
    "ee": "Ewe",
    "fo": "Faroese",
    "fj": "Fijian",
    "fil": "Filipino",
    "fi": "Finnish",
    "fr": "French",
    "gaa": "Ga",
    "gl": "Galician",
    "lg": "Ganda",
    "ka": "Georgian",
    "de": "German",
    "el": "Greek",
    "gn": "Guarani",
    "gu": "Gujarati",
    "ht": "Haitian Creole",
    "ha": "Hausa",
    "haw": "Hawaiian",
    "iw": "Hebrew",
    "hi": "Hindi",
    "hmn": "Hmong",
    "hu": "Hungarian",
    "is": "Icelandic",
    "ig": "Igbo",
    "id": "Indonesian",
    "ga": "Irish",
    "it": "Italian",
    "ja": "Japanese",
    "jv": "Javanese",
    "kl": "Kalaallisut",
    "kn": "Kannada",
    "kk": "Kazakh",
    "kha": "Khasi",
    "km": "Khmer",
    "rw": "Kinyarwanda",
    "ko": "Korean",
    "kri": "Krio",
    "ku": "Kurdish",
    "ky": "Kyrgyz",
    "lo": "Lao",
    "la": "Latin",
    "lv": "Latvian",
    "ln": "Lingala",
    "lt": "Lithuanian",
    "luo": "Luo",
    "lb": "Luxembourgish",
    "mk": "Macedonian",
    "mg": "Malagasy",
    "ms": "Malay",
    "ml": "Malayalam",
    "mt": "Maltese",
    "gv": "Manx",
    "mi": "Māori",
    "mr": "Marathi",
    "mn": "Mongolian",
    "mfe": "Morisyen",
    "ne": "Nepali",
    "new": "Newari",
    "nso": "Northern Sotho",
    "no": "Norwegian",
    "ny": "Nyanja",
    "oc": "Occitan",
    "or": "Odia",
    "om": "Oromo",
    "os": "Ossetic",
    "pam": "Pampanga",
    "ps": "Pashto",
    "fa": "Persian",
    "pl": "Polish",
    "pt": "Portuguese",
    "pt-PT": "Portuguese (Portugal)",
    "pa": "Punjabi",
    "qu": "Quechua",
    "ro": "Romanian",
    "rn": "Rundi",
    "ru": "Russian",
    "sm": "Samoan",
    "sg": "Sango",
    "sa": "Sanskrit",
    "gd": "Scottish Gaelic",
    "sr": "Serbian",
    "crs": "Seselwa Creole French",
    "sn": "Shona",
    "sd": "Sindhi",
    "si": "Sinhala",
    "sk": "Slovak",
    "sl": "Slovenian",
    "so": "Somali",
    "st": "Southern Sotho",
    "es": "Spanish",
    "su": "Sundanese",
    "sw": "Swahili",
    "ss": "Swati",
    "sv": "Swedish",
    "tg": "Tajik",
    "ta": "Tamil",
    "tt": "Tatar",
    "te": "Telugu",
    "th": "Thai",
    "bo": "Tibetan",
    "ti": "Tigrinya",
    "to": "Tongan",
    "ts": "Tsonga",
    "tn": "Tswana",
    "tum": "Tumbuka",
    "tr": "Turkish",
    "tk": "Turkmen",
    "uk": "Ukrainian",
    "ur": "Urdu",
    "ug": "Uyghur",
    "uz": "Uzbek",
    "ve": "Venda",
    "vi": "Vietnamese",
    "war": "Waray",
    "cy": "Welsh",
    "fy": "Western Frisian",
    "wo": "Wolof",
    "xh": "Xhosa",
    "yi": "Yiddish",
    "yo": "Yoruba",
    "zu": "Zulu"
}

def markdown_result(result):
    return display(Markdown(result))
def tokenize(lang):
    return word_tokenize(lang)
from langchain.document_loaders import YoutubeLoader
def youtubetranscript(links,language=None,translation='en'):
    text=''
    if language=="en":
    
        for i in links.split(','):
            loader=YoutubeLoader.from_youtube_url(i,add_video_info=True)
            Transcript=loader.load()
            text+=Transcript[0].page_content
        return text
    else:
        for i in links.split(','):
            loader=YoutubeLoader.from_youtube_url(i,language=language,translation=translation,add_video_info=True)
            Transcript=loader.load()
            text+=Transcript[0].page_content
        return text


    

def videogenre(models,transcript):
    
    prompt = '''
    Based on the provided YouTube video transcript, classify the genre into one of the following categories:
    - Case Study
    - Educational/Tutorial
    - Documentary
    - Podcast/Interview/Debate
    - News/Current Affairs
    - Other

    Please provide only the genre from the list above.

    Input Transcript:
    {Transcript}
    '''
    template = PromptTemplate(input_variables=['Transcript'], template=prompt)
    content_genre=models.invoke(template.format(Transcript=transcript)).content
    return content_genre
def web_data(link):
    text=''
    for i in link.split(','):
        loader=WebBaseLoader(i)
        text+=loader.load()
    return text
def add_html_to_docx(html, doc):
    for element in html:
        if element.name == "h1":
            run = doc.add_heading(level=1).add_run(element.text)
            run.font.size = Pt(24)
        elif element.name == "p":
            p = doc.add_paragraph()
            run = p.add_run(element.text)
        elif element.name == "ul":
            for li in element.find_all("li"):
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(li.text)
def split_text(text):
    text_split=RecursiveCharacterTextSplitter(chunk_size=4000,chunk_overlap=500)
    return text_split.split_text(text)
def qna_chain(models,text,question):
    chain=load_qa_chain(models,chain_type='map_rerank',verbose=True,return_intermediate_steps=True)
    result=chain({'input_documents':text,'question':question},return_only_outputs=True)
    return result['output_text']

def report_creation(models, text, format, type_of_report):
    report_template = """You have to generate a {type_of_report} report based on the following data:
    '{text}'
    The format for the report is as follows:
    '{format}'"""
    
    # Create the prompt
    prompt = PromptTemplate(template=report_template, input_variables=['text', 'format', 'type_of_report'])
    
    # Load the summarize chain without passing the prompt directly
    result=models.invoke(prompt.format(text=text,format=format,type_of_report=type_of_report))
    
    # Call the chain with the required inputs
    
    return result

def summarize_video(Transcript, models, content_genre):
    content_genre = content_genre.replace('\n', '').strip()
    
    if content_genre == 'Educational/Tutorial':
        prompt = f'''Summarize this transcript in points and briefly describe it. If the video involves coding, include the piece of code as well. from the following:
        {Transcript} Replace the word Transcript with Video in result'''
    
    elif content_genre in ['Case Study', 'News/Current Affairs']:
        prompt = f''' Summarize this transcript in detail, describing each event in the order of occurrence.from the following:
        {Transcript} Replace the word Transcript with Video in result'''
        
    elif content_genre in ['Podcast/Interview/Debate', 'Documentary']:
        prompt = f'''Explain this transcript in detail with subheadings and points from the following :
        {Transcript} Replace the word Transcript with Video in result'''
        
    else:
        prompt = f''' Provide a general summary of this transcript. from the following:
        {Transcript} Replace the word Transcript with Video in result'''
    
    print(prompt)
    template = PromptTemplate(input_variables=['Transcript'], template=prompt)
    result = models.invoke(template.format(Transcript=Transcript)).content
    return result


def summarize_web_and_vid(Text):
    prompt="""Write the concise summary of the following:
    "{text}" 
    CONCISE SUMMARY:
    """
    combine_prompt="""
    Write a concise summary of the following text delimited by triple backqoutes,
    Return your response in bullet points which covers the key points of the text.
    ```{text}```
    BULLET POINT SUMMARY:
    """
    combine_prompt_template=PromptTemplate(template=combine_prompt_template,input_variables=['text'])
    summary_prompt=PromptTemplate(template=prompt,input_variables=['text'])
    summary_chain=load_summarize_chain(llm=models,chain_type='map_reduce',map_prompt=summary_prompt,combine_prompt=combine_prompt_template)
    output=summary_chain(Text)
    return output
def match_back_language(value):
    for key,values in languages.items():
        if values==value:
            return key






